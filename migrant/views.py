from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from django.template.loader import get_template
from rest_framework.views import APIView
from rest_framework.response import Response
from .forms import CaseForm, CompanyForm, IndividualForm, GroupForm, EntrepreneurForm, PhotoForm, FileForm, CaseCommentForm
from .filters import MigrantFilter
from .serializers import *
# Create your views here.
from .models import *
from main.service import unpucking
from django.db import connection
from django.http import Http404
from docxtpl import DocxTemplate
from .templatetags.migrant_tags import var_verbose_name_for_word, check_arg_is_none
from docx import Document
import jinja2
import pdfkit
import os



@login_required
def append_case(request):
    if request.method == 'POST':
        form = CaseForm(request.POST)
        companyForm = CompanyForm(request.POST)
        individualForm = IndividualForm(request.POST)
        groupForm = GroupForm(request.POST)
        entrepreneurForm = EntrepreneurForm(request.POST)
        photoForm = PhotoForm(request.POST, request.FILES)
        fileForm = FileForm(request.POST, request.FILES)

        if form.is_valid():

            # cleaned_data = form.cleaned_data;
            # if request.user.is_authenticated():
            #     form.instance.user = request.user
            # if cleaned_data['source']:
            #     for item in cleaned_data['source']:
            #         if item.id == 5:
            #             pass
            case = form.save(commit=False)


            if companyForm.is_valid():
                case.company = companyForm.save()
            if individualForm.is_valid():
                case.individualInfo = individualForm.save()
            # if victimForm.is_valid():
            #     victim = victimForm.save(commit=False)
            #     case.victim = victim
            #     victim.save()
            if groupForm.is_valid():
                case.personGroupInfo = groupForm.save()
            if entrepreneurForm.is_valid():
                case.entrepreneur = entrepreneurForm.save()

            case.user = request.user
            case.save()
            form._save_m2m()

            if photoForm.is_valid():
                for f in request.FILES.getlist('photo'):
                    photo = CasePhoto(photo=f, card=case)
                    photo.save()
            if fileForm.is_valid():
                for f in request.FILES.getlist('file'):
                    file = CaseFile(file=f, card=case)
                    file.save()

            return redirect('migrants_list')

    else:
        form = CaseForm
        companyForm = CompanyForm
        individualForm = IndividualForm
        groupForm = GroupForm
        entrepreneurForm = EntrepreneurForm
        photoForm = PhotoForm
        fileForm = FileForm

    return render(request,
                  'migrant/add_case.html',
                  context={
                      'form': form,
                      'companyForm': companyForm,
                      'individualForm': individualForm,
                      'groupForm': groupForm,
                      'entrepreneurForm': entrepreneurForm,
                      'photoForm': photoForm,
                      'fileForm': fileForm,
                  })
@login_required()
def delete_case(request, pk):
    case = Case.objects.get(id=pk).delete()
    return redirect('migrants_list')

@login_required()
def update_case(request,pk):
    case = Case.objects.get(id=pk)

    if request.method == 'POST':
        form = CaseForm(request.POST, instance=case)
        companyForm = CompanyForm(request.POST)
        individualForm = IndividualForm(request.POST)
        groupForm = GroupForm(request.POST)
        entrepreneurForm = EntrepreneurForm(request.POST)
        photoForm = PhotoForm(request.POST, request.FILES)
        fileForm = FileForm(request.POST, request.FILES)

        if form.is_valid():
            form.save(commit=False)

        if companyForm.is_valid():
            case.company = companyForm.save()
        if individualForm.is_valid():
            case.individualInfo = individualForm.save()
        if groupForm.is_valid():
            case.personGroupInfo = groupForm.save()
        if entrepreneurForm.is_valid():
            case.entrepreneur = entrepreneurForm.save()

        case.user = request.user
        case.save()
        form._save_m2m()

        if photoForm.is_valid():
            for f in request.FILES.getlist('photo'):
                photo = CasePhoto(photo=f, card=case)
                photo.save()
        if fileForm.is_valid():
            for f in request.FILES.getlist('file'):
                file = CaseFile(file=f, card=case)
                file.save()

        return redirect('migrants_list')
    else:
        form = CaseForm(instance=case)

        companyForm = CompanyForm()
        if case.company is not None:
            companyForm = CompanyForm(instance=Company.objects.get(id=case.company))

        individualForm = IndividualForm
        if case.individualInfo is not None:
            individualForm = IndividualForm(instance=IndividualInfo.objects.get(id=case.individualInfo_id))

        # victimForm = VictimForm
        # if case.victim is not None:
        #     victimForm = VictimForm(instance=Victim.objects.get(id=case.victim_id))

        groupForm = GroupForm()
        if case.personGroupInfo is not None:
            groupForm = GroupForm(instance=PersonGroup.objects.get(id = case.personGroupInfo_id))

        entrepreneurForm = EntrepreneurForm
        if case.entrepreneur is not None:
            entrepreneurForm = EntrepreneurForm(instance = Entrepreneur.objects.get(id = case.entrepreneur_id))
        #
        # photos = CasePhoto.objects.get(card = case.id)
        # print(photos)
        photoForm = PhotoForm
        fileForm = FileForm
        images = CasePhoto.objects.filter(card_id=case.id)
        files = CaseFile.objects.filter(card_id=case.id)
    return render(request,
                  'migrant/add_case.html',
                  context={
                      'form': form,
                      'companyForm': companyForm,
                      'individualForm': individualForm,
                      # 'victimForm': victimForm,
                      'groupForm': groupForm,
                      'entrepreneurForm': entrepreneurForm,
                      'photoForm': photoForm,
                      'fileForm': fileForm,
                      'images': images,
                      'files': files,
                  })

@login_required()
def cases(request):
    if request.user.position.role_id == 1:
        cases = Case.objects.all()
        filter = MigrantFilter(request.GET, queryset=cases)
        cards = filter.qs
        context = {'cards': cards, 'myFilter': filter}
        return render(request, 'migrant/cases.html', context)
    elif request.user.position.role_id == 2 or request.user.position.role_id == 3:
        cases = Case.objects.filter(country_id=request.user.country.country_id)
        filter = MigrantFilter(request.GET, queryset=cases)
        cards = filter.qs
        context = {'cards': cards, 'myFilter': filter}
        return render(request, 'migrant/cases.html', context)
    else:
        raise Http404('Недостаточно прав!')

def multistep(request):
    return render(request, 'migrant/add_case.html')


def download(request, case_id):
    pass


def load_regions(request):
    country_id = request.GET.get('country_id')
    regions = Region.objects.filter(country=country_id)
    return render(request,'migrant/region_dropdown.html',{'regions':regions})

@login_required()
def add_comment(request, pk):
    case = Case.objects.get(id=pk)
    if request.method == 'POST':
        data = request.POST
        data._mutable = True
        data['case'] = case.id
        form = CaseCommentForm(data)
        if form.is_valid():
            form.save()
            return redirect('migrants_list')
    else:
        form = CaseCommentForm()
    return render(request, 'migrant/add_comment.html', {'form': form, 'case': case})


def show_comments(request, pk):
    comments = CaseComment.objects.filter(case_id=pk, active=True)
    return render(request, 'migrant/show_comments.html', {'comments': comments})


def delete_comment(request, pk):
    CaseComment.objects.get(id=pk).delete()
    return redirect('migrants_list')


def case_render_pdf_view(request, *args, **kwargs):
    pk = kwargs.get('pk')
    case = get_object_or_404(Case, pk=pk)
    source = InfoSource.objects.filter(case__pk=pk)
    right = Right.objects.filter(case__pk=pk)
    intruder = Intruder.objects.filter(case__pk=pk)
    comments = CaseComment.objects.filter(case_id=pk)
    template_path = 'migrant/migrant_pdf.html'
    context = {
        'case': case,
        'source': source,
        'right': right,
        'intruder': intruder,
        'comments': comments,
    }
    # Create a Django response object, and specify content_type as pdf
    # find the template and render it.
    template = get_template(template_path)
    html = template.render(context)
    # print(html)
    # create a pdf
    pdf = pdfkit.from_string(html, False)
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'filename="card_{case.id}.pdf"'
    # if error then show some funy view
    # if pisa_status.err:
    #     return HttpResponse('We had some errors <pre>' + html + '</pre>')
    return response


def case_download_pdf_view(request, *args, **kwargs):
    pk = kwargs.get('pk')
    case = get_object_or_404(Case, pk=pk)
    source = InfoSource.objects.filter(case__pk=pk)
    right = Right.objects.filter(case__pk=pk)
    intruder = Intruder.objects.filter(case__pk=pk)
    comments = CaseComment.objects.filter(case_id=pk)
    template_path = 'migrant/migrant_pdf.html'
    context = {
        'case': case,
        'source': source,
        'right': right,
        'intruder': intruder,
        'comments': comments,
    }
    # Create a Django response object, and specify content_type as pdf
    # find the template and render it.
    template = get_template(template_path)
    html = template.render(context)
    # create a pdf
    pdf = pdfkit.from_string(html, False)
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="card_{case.id}.pdf"'
    # if error then show some funy view
    # if pisa_status.err:
    #     return HttpResponse('We had some errors <pre>' + html + '</pre>')
    return response

def test(request, pk):
    case = Case.objects.get(pk=pk)
    photos = CasePhoto.objects.filter(card_id=pk)
    print(photos)
    for i in photos:
        print(i)
    return render(request, 'migrant/migrant_pdf.html', {'case': case, 'photos': photos})


class DataAPIView(APIView):
    def get(self, request):
        country = CountrySerializers(Country.objects.all(), many=True)
        region_1 = RegionSerializers(Region.objects.filter(country__pk=1), many=True)
        region_2 = RegionSerializers(Region.objects.filter(country__pk=2), many=True)
        region_3 = RegionSerializers(Region.objects.filter(country__pk=3), many=True)
        region_4 = RegionSerializers(Region.objects.filter(country__pk=4), many=True)
        region_5 = RegionSerializers(Region.objects.filter(country__pk=5), many=True)
        victim = VictimSerializers(Victim.objects.all(), many=True)
        banOnEntry = BanOnEntrySerializers(BanOnEntry.objects.all(), many=True)
        source = InfoSourceSerializers(InfoSource.objects.all(), many=True)
        violated_right = RightSerializers(Right.objects.all(), many=True)
        victim = VictimSerializers(Victim.objects.all(), many=True)
        # individualInfo = IndividualInfoSerializers(IndividualInfo.objects.all(), many=True)
        intruder = IntruderSerializers(Intruder.objects.all(), many=True)
        violation_nature = NatureViolationSerializers(NatureViolation.objects.all(), many=True)
        # personGroupInfo = PersonGroupSerializers(PersonGroup.objects.all(), many=True)
        entrepreneur = EntrepreneurSerializers(Entrepreneur.objects.all(), many=True)
        rights_state = RightsStateSerializers(RightsState.objects.all(), many=True)
        victim_situation = VictimSituationSerializers(VictimSituation.objects.all(), many=True)
        tradeUnionSituation = TradeUnionSituationSerializers(TradeUnionSituation.objects.all(), many=True)
        violationType = ViolationTypeSerializers(ViolationType.objects.all(), many=True)
        changesInSalary = ChangesInSalarySerializers(ChangesInSalary.objects.all(), many=True)
        user = UserSerializers(User.objects.all(), many=True)
        # company = CompanySerializers(Company.objects.all(), many=True)
        # tradeUnionCount = TradeUnionCountSerializers(TradeUnionCount.objects.all(), many=True)
        return Response([
            {'id': 'country', 'name': 'Страна', 'item': country.data},
            {'id': 'region', 'name': 'Регион', 'item': {1: region_1.data,
                                                        2: region_2.data,
                                                        3: region_3.data,
                                                        4: region_4.data,
                                                        5: region_5.data}},
            {'id': 'victim', 'name': 'В отношении кого совершено нарушение', 'item': victim.data},
            {'id': 'banonentry', 'name': 'Есть ли у вас запрет на въезд?', 'item': banOnEntry.data},
            {'id': 'infosource', 'name': 'Источник информации о нарушении', 'item': source.data},
            {'id': 'right', 'name': 'Какое право нарушено?', 'item': violated_right.data},
            # {'id': 'individualinfo', 'name': 'физическое лицо', 'item': individualInfo.data}, #50 % 50
            # {'id': 'personGroupInfo', 'name': 'Группа лиц', 'item': personGroupInfo.data},
            {'id': 'entrepreneur', 'name': 'Работодатель(Частное лицо)', 'item': entrepreneur.data},
            {'id': 'intruder', 'name': 'Кем было совершено нарушение', 'item': intruder.data},
            {'id': 'natureviolation', 'name': 'Характер нарушения', 'item': violation_nature.data},
            {'id': 'rightsstate', 'name': 'Ситуация с правами', 'item': rights_state.data},
            {'id': 'victimsituation', 'name': 'Ситуация с потерпевшим(и)', 'item': victim_situation.data},
            {'id': 'tradeunionsituation', 'name': 'Профсоюз на месте работы после произошедшего',
             'item': tradeUnionSituation.data},
            {'id': 'violationtype', 'name': 'С какими нарушениями трудовых прав вы столкнулись из-за COVID-19?',
             'item': violationType.data},
            {'id': 'changesinsalary', 'name': 'Как изменились Ваши доходы из-за COVID-19?', 'item': changesInSalary.data},
            {'id': 'user', 'name': 'Монитор', 'item': user.data},
        ])
        # {'id': 'company', 'name': 'Работодатель(компания)', 'item': company.data},  # Можно удалить
        # {'id': 'tradeUnionCount', 'name': 'Численность профсоюза после произошедшего',
        #  'item': tradeUnionCount.data},  # Можно удалить

class DataFilterAPI(APIView):
    authentication_classes = []
    def post(self, request):
        my_list = []
        print(request.data)
        for item in request.data:
            if item['id'] == 'user':
                my_list.append(f"auth_user.username")
            elif item['id'] == 'entrepreneur':
                my_list.append(f"migrant_entrepreneur.entrepreneur_name")
            else:
                my_list.append(f"migrant_{item['id']}.name")
        fields = unpucking(my_list)
        # print(my_list)
        case_count = Case.objects.count()
        sql_query = f"SELECT {fields}, count(*), count(*), 100 / count (*) FROM migrant_case"
        where_query = "where "
        where_list = []
        where_query_list = []
        group_by_query = f"group by {fields}"
        for data in request.data:
            if data['id'] in fields:
                id = data['id']
                item = data['item']
                if id == "country":
                    where_sql_query = "migrant_case.country_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)} ")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_country on migrant_country.id = migrant_case.country_id "

                elif id == "region":
                    where_sql_query = "migrant_case.region_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_region on migrant_region.id = migrant_case.region_id "

                elif id == "banonentry":
                    where_sql_query = "migrant_case.banonentry_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_banonentry on migrant_banonentry.id = migrant_case.banonentry_id "

                elif id == "individualinfo":
                    where_sql_query = "migrant_case.individualinfo_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_individualinfo on migrant_individualinfo.id = migrant_case.individualinfo_id "

                elif id == "entrepreneur":
                    where_sql_query = "migrant_case.entrepreneur_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_entrepreneur on migrant_entrepreneur.id = migrant_case.entrepreneur_id "

                elif id == "violationtype":
                    where_sql_query = "migrant_case.violationtype_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_violationtype on migrant_violationtype.id = migrant_case.violationtype_id "

                elif id == "changesinsalary":
                    where_sql_query = "migrant_case.changesinsalary_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_changesinsalary on migrant_changesinsalary.id = migrant_case.changesinsalary_id "

                elif id == "victim":
                    where_sql_query = "migrant_case.victim_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_victim on migrant_victim.id = migrant_case.victim_id "

                elif id == "natureviolation":
                    where_sql_query = "migrant_case.violation_nature_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_natureviolation on migrant_natureviolation.id = migrant_case.violation_nature_id "

                elif id == "rightsstate":
                    where_sql_query = "migrant_case.rights_state_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_rightsstate on migrant_rightsstate.id = migrant_case.rights_state_id "

                elif id == "victimsituation":
                    where_sql_query = "migrant_case.victim_situation_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_victimsituation on migrant_victimsituation.id = migrant_case.victim_situation_id "

                elif id == "tradeunionsituation":
                    where_sql_query = "migrant_case.tradeunionsituation_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_tradeunionsituation on migrant_tradeunionsituation.id = migrant_case.tradeunionsituation_id "

                elif id == "user":
                    where_sql_query = "migrant_case.user_id in "
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join auth_user on auth_user.id = migrant_case.user_id "

                # Запрос для нахождение между start_date and end_date
                elif id == 'start_date':
                    # Переменные где будут находится даты
                    start_date, end_date = 1, 2
                    # Сам запрос
                    where_query_list.append(f"work_case.start_date BETWEEN date({item[0]}) AND date({item[1]})")


                # elif id == "": # Экземпляр
                #     where_sql_query = "migrant_case._id in "
                #     for i in item:
                #         where_list.append(i['id'])
                #     if len(where_list) > 1:
                #         where_query_list.append(f"{where_sql_query} {tuple(where_list)}")
                #     else:
                #         where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                #     where_list.clear()
                #     sql_query += " join migrant_ on migrant_.id = migrant_case._id "


                # Ниже представлены ManyToMany связи!
                elif id == "infosource":
                    where_sql_query = "migrant_case_source.infosource_id in"
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)} ")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_case_source on migrant_case.id = migrant_case_source.case_id join migrant_infosource on migrant_case_source.infosource_id = migrant_infosource.id "

                elif id == "intruder":
                    where_sql_query = "migrant_case_intruder.intruder_id in"
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)} ")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_case_intruder on migrant_case.id = migrant_case_intruder.case_id join migrant_intruder on migrant_case_intruder.intruder_id = migrant_intruder.id "

                elif id == "right":
                    where_sql_query = "migrant_case_violated_right.right_id in"
                    for i in item:
                        where_list.append(i['id'])
                    if len(where_list) > 1:
                        where_query_list.append(f"{where_sql_query} {tuple(where_list)} ")
                    else:
                        where_query_list.append(f"{where_sql_query} ({where_list[0]}) ")
                    where_list.clear()
                    sql_query += " join migrant_case_violated_right on migrant_case.id = migrant_case_violated_right.case_id join migrant_right on migrant_case_violated_right.right_id = migrant_right.id "
            else:
                continue
        where_query_list = 'and '.join(where_query_list)
        where_query += where_query_list
        # print(where_query)
        # print(sql_query + where_query + group_by_query)
        with connection.cursor() as cursor:
            cursor.execute(
                sql_query + where_query + group_by_query
            )
            row = cursor.fetchall()
            print(row)
            fields_list = []
            for i in request.data:
                fields_list.append(i['id'])
            fields_list.append('count')
            fields_list.append('percent')
            response_list = []

            for i in range(len(row)):
                response_body = dict()
                for j in range(len(fields_list)):
                    response_body[fields_list[j]] = row[i][j]
                response_list.append(response_body)
            for i in range(len(response_list)):
                response_list[i]['percent'] = 100 / len(response_list)
        return Response(response_list)

@login_required()
def case_files_download(request, pk):
    if request.user.position.role_id == 3:
        case = Case.objects.get(user=request.user, pk=pk)
        images = CasePhoto.objects.filter(card_id=case.id)
        files = CaseFile.objects.filter(card_id=case.id)
        return render(request, 'migrant/files_download.html', {'images': images, 'files': files})
    else:
        case = Case.objects.get(pk=pk)
        images = CasePhoto.objects.filter(card_id=case.id)
        files = CaseFile.objects.filter(card_id=case.id)
        return render(request, 'migrant/files_download.html', {'images': images, 'files': files})


@login_required()
def case_files_delete(request, pk):
    case_file = CaseFile.objects.get(pk=pk)
    case_id = case_file.card_id
    case_file.delete()
    return redirect('migrant_case_files_download', pk=case_id)


@login_required()
def case_photo_delete(request, pk):
    case_photo = CasePhoto.objects.get(pk=pk)
    case_id = case_photo.card_id
    case_photo.delete()
    return redirect('migrant_case_files_download', pk=case_id)


def migrant_word_generate(request, pk):
    case = Case.objects.get(pk=pk)
    sources = InfoSource.objects.filter(case__pk=pk)
    rights = Right.objects.filter(case__pk=pk)
    intruders = Intruder.objects.filter(case__pk=pk)
    document = Document()
    dates_list = ['date_create', 'date_update', 'start_date', 'end_date']
    document.add_heading('Мигранты', 0)
    records = []
    case_values = Case.objects.filter(pk=pk).values()
    field_name_list = [i.name for i in Case._meta.get_fields()]
    for field in field_name_list:
        try:
            if field == 'source':
                verbose_name = Case._meta.get_field(field).verbose_name
                for source in sources:
                    source_name = source.name
                    if source_name is not None and source_name != '':
                        records.append((verbose_name, source_name))
            elif field == 'violated_right':
                verbose_name = Case._meta.get_field(field).verbose_name
                for right in rights:
                    right_name = right.name
                    if right_name is not None and right_name != '':
                        records.append((verbose_name, right_name))
            elif field == 'intruder':
                verbose_name = Case._meta.get_field(field).verbose_name
                for intruder in intruders:
                    intruder_name = intruder.name
                    if intruder_name is not None and intruder_name != '':
                        records.append((verbose_name, intruder_name))
            elif field == 'casephoto':
                continue
            elif field == 'casefile':
                continue
            elif field == 'casecomment':
                continue
            elif field == 'active':
                continue
            elif field == 'id':
                continue
            else:
                verbose_name = Case._meta.get_field(field).verbose_name
                value = case_values[0][field]
                if value is not None and value != '':
                    if field in dates_list:
                        records.append((verbose_name, value.strftime("%Y.%m.%d")))
                    else:
                        records.append((verbose_name, value))
        except KeyError:
            try:
                field += '_id'
                value_id = case_values[0][field]
                if value_id is not None and value_id != '':
                    value_name = Case._meta.get_field(field).related_model.objects.get(pk=value_id)
                    if field == 'tradeUnionCount_id':
                        records.append((verbose_name, value_name.choice))
                    elif field == 'personGroupInfo_id':
                        records.append((verbose_name, value_name.amount))
                    elif field == 'user_id':
                        records.append((verbose_name, value_name.username))
                    else:
                        records.append((verbose_name, value_name.name))
                else:
                    continue
            except Exception as e:
                print(e, field)

    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for field, value in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(field)
        row_cells[1].text = str(value)

    document.add_page_break()
    base_dir = str(settings.BASE_DIR)
    base_dir += "/migrant/static/word/migrant/"
    save_path = base_dir + f'card_{case.id}.docx'
    document.save(save_path)
    if os.path.exists(save_path):
        with open(save_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(save_path)
            return response
    return Http404()
