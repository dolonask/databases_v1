from django.contrib.auth.decorators import login_required
from django.core.mail import send_mail
from django.db.models import Count
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, Http404
from django.template.loader import get_template
from rest_framework.views import APIView
from rest_framework.response import Response

from strike.helpers import get_request_data, return_right_data, get_values, get_data, get_fields
from .functions import remove_data
from .serializers import *
from django.contrib.auth.models import User
from django.db import connection
from main.service import unpucking
from .filters import WorkFilter
from .forms import CaseForm, IndividualForm, PersonGroupForm, TradeUnionInfoForm, CompanyInfoForm, CasePhotoForm, \
    CaseFileForm, IndividualFormSet, CaseCommentForm
from .models import *
import os
import pdfkit
from docx import Document

# Create your views here.

general_tabs_fields = [
    'case_name',
    'source',
    'source_another',
    'source_url',
    'source_content',
    'country',
    'region',
    'city_name',
    'trade_union_activities',
    'trade_union_activities_another',
    'case_company_name',
    'groupOfRights',
    'tradeUnionRight',
    'tradeUnionRightAnother',
    'tradeUnionCrime',
    'tradeUnionCrimeAnother',
    'meetingsRight',
    'meetingsRightAnother',
    'сonvention87',
    'сonvention98',
    'conversationRight',
    'conversationRightAnother',
    'principleOfNonDiscrimination',
    'discriminationInVariousAreas',
    'discriminationInVariousAreasAnother',
    'tradeUnionBuildingsRight',
    'tradeUnionBuildingsRightAnother',
    'createOrganizationRight',
    'createOrganizationRightAnother',
    'createTradeUnionRight',
    'createTradeUnionRightAnother',
    'electionsRight',
    'electionsRightAnother',
    'tradeUnionActivityRight',
    'tradeUnionActivityRightAnother',
    'createStrikeRight',
    'createStrikeRightAnother',
    'antiTradeUnionDiscrimination',
    'antiTradeUnionDiscriminationAnother',
    'сonvention135',
    'сonvention135Another',
    'consultationRight',
    'consultationRightAnother',
    'discriminatiOnVariousGrounds',
    'publicPolicyDiscrimination',
    'childLabor',
    'сonvention138',
    'convention182',
    'prohibitionOfForcedLabor',
    'useOfForcedLabor',
    'governmentCoercion',
    'violationsUsingCompulsoryLabor',
    'failureSystemicMeasures',
    # 'date_type',
    'start_date',
    'end_date',
]
initiator_tab_fields = [
    'victim'
]
intruder_tab_fields = [
    'intruder',
    'intruderAnother',
    'government_agency_name',
    'local_agency_name',
    'police_agency_name',
    # 'control_agency_name',
]
description_tab_fields = [
    'exact_data',
    'case_description',
    'tradeunion_actions',
    'case_result',
    'violation_nature',
    'rights_state',
    'rights_state_another',
    'victim_situation',
    'victim_situation_another',
    'tradeUnionSituation',
    'tradeUnionSituation_another',
    'tradeUnionCount',
]
files_tab_fields = [
    'case_text'
]

@login_required()
def add_case(request):
    if request.method == 'POST':
        form = CaseForm(request.POST)
        tradeUnionForm = TradeUnionInfoForm(request.POST)
        individualForm = IndividualForm(request.POST)
        individualFormSet = IndividualFormSet(data=request.POST)
        personGroupForm = PersonGroupForm(request.POST)
        companyInfoForm = CompanyInfoForm(request.POST)
        casePhotoForm = CasePhotoForm(request.POST)
        caseFileForm = CaseFileForm(request.POST)

        if form.is_valid():
            case = form.save(commit=False)

            if tradeUnionForm.is_valid():
                case.tradeUnionInfo = tradeUnionForm.save()
            if personGroupForm.is_valid():
                case.groupOfPersons = personGroupForm.save()
            if companyInfoForm.is_valid():
                case.company = companyInfoForm.save()


            case.user = request.user
            case.active = True
            case.save()

            if individualFormSet.is_valid():
                individualFormSet.save()

            form._save_m2m()

            for individual in individualFormSet:
                if individual.is_valid():
                    ind = individual.save(commit=False)
                    ind.case_id = case.id
                    ind.save()

            # if casePhotoForm.is_valid():
            for f in request.FILES.getlist('photo'):
                photo = CasePhoto(photo=f, card=case)
                photo.save()
            # if caseFileForm.is_valid():
            for f in request.FILES.getlist('file'):
                file = CaseFile(file=f, card=case)
                file.save()

            return redirect('works_list')
    else:
        form = CaseForm
        tradeUnionForm = TradeUnionInfoForm
        individualForm = IndividualForm
        individualFormSet = IndividualFormSet(queryset=IndividualInfo.objects.none())
        personGroupForm = PersonGroupForm
        companyInfoForm = CompanyInfoForm
        casePhotoForm = CasePhotoForm
        caseFileForm = CaseFileForm
    return render(request, 'work/add_case.html', context={
        'form': form,
        'tradeUnionForm': tradeUnionForm,
        'individualFormSet': individualFormSet,
        'personGroupForm': personGroupForm,
        'companyInfoForm': companyInfoForm,
        'caseFileForm': caseFileForm,
        'casePhotoForm': casePhotoForm,
        'general_tabs_fields': general_tabs_fields,
        'initiator_tab_fields': initiator_tab_fields,
        'intruder_tab_fields': intruder_tab_fields,
        'description_tab_fields': description_tab_fields,
        'files_tab_fields': files_tab_fields,
        'individualForm': individualForm,
    })


@login_required()
def update_case(request, pk):
    case = Case.objects.get(id=pk)
    if request.method == 'POST':
        form = CaseForm(request.POST, instance=case)
        tradeUnionForm = TradeUnionInfoForm(request.POST)
        individualForm = IndividualForm(request.POST)
        individualFormSet = IndividualFormSet(data=request.POST)
        personGroupForm = PersonGroupForm(request.POST)
        companyInfoForm = CompanyInfoForm(request.POST)
        casePhotoForm = CasePhotoForm(request.POST)
        caseFileForm = CaseFileForm(request.POST)

        if form.is_valid():
            form.save(commit=False)

            for individual in individualFormSet:
                if individual.is_valid():
                    ind = individual.save(commit=False)
                    ind.case = case
                    ind.save()

            if tradeUnionForm.is_valid():
                case.tradeUnionInfo = tradeUnionForm.save()
            if personGroupForm.is_valid():
                case.groupOfPersons = personGroupForm.save()
            if companyInfoForm.is_valid():
                case.company = companyInfoForm.save()


            case.user = request.user
            case.active = True
            case.save()

            if individualFormSet.is_valid():
                individualFormSet.save()

            form._save_m2m()


            # if casePhotoForm.is_valid():
            for f in request.FILES.getlist('photo'):
                photo = CasePhoto(photo=f, card=case)
                photo.save()
            # if caseFileForm.is_valid():
            for f in request.FILES.getlist('file'):
                file = CaseFile(file=f, card=case)
                file.save()

            return redirect('works_list')
    form = CaseForm(instance=case)

    individualForm = IndividualForm

    tradeUnionForm = TradeUnionInfoForm
    if case.tradeUnionInfo is not None:
        tradeUnionForm = TradeUnionInfoForm(instance=TradeUnionInfo.objects.get(pk=case.tradeUnionInfo_id))

    personGroupForm = PersonGroupForm
    if case.groupOfPersons is not None:
        personGroupForm = PersonGroupForm(instance=GroupOfPersons.objects.get(pk=case.groupOfPersons_id))


    companyInfoForm = CompanyInfoForm

    if case.company is not None:
        companyInfoForm = CompanyInfoForm(instance=Company.objects.get(pk=case.company_id))

    casePhotoForm = CasePhotoForm()
    caseFileForm = CaseFileForm()
    individual_infos = IndividualInfo.objects.filter(case__id=pk)
    images = CasePhoto.objects.filter(card_id=case.id)
    files = CaseFile.objects.filter(card_id=case.id)
    individualFormSet = IndividualFormSet(queryset=IndividualInfo.objects.filter(case_id=case.id))
    if len(individual_infos) != 0:
        individualFormSet.extra = 0
    return render(request, 'work/add_case.html', context={
        'form':form,
        'tradeUnionForm':tradeUnionForm,
        'individualFormSet':individualFormSet,
        'personGroupForm':personGroupForm,
        'companyInfoForm':companyInfoForm,
        'caseFileForm':caseFileForm,
        'casePhotoForm':casePhotoForm,
        'general_tabs_fields':general_tabs_fields,
        'initiator_tab_fields':initiator_tab_fields,
        'intruder_tab_fields':intruder_tab_fields,
        'description_tab_fields':description_tab_fields,
        'files_tab_fields':files_tab_fields,
        'images': images,
        'files': files,
        'individual_infos': individual_infos
    })


def delete_case(request, pk):
    case = Case.objects.get(id=pk).delete()
    return redirect('works_list')


@login_required()
def cases(request):
    if request.user.position.role_id == 1:
        cases = Case.objects.all()
        filter = WorkFilter(request.GET, queryset=cases)
        cards = filter.qs
        context = {'cards': cards, 'myFilter': filter}
        return render(request, 'work/cases.html', context)
    elif request.user.position.role_id == 2:
        cases = Case.objects.filter(country_id=request.user.country.country_id)
        filter = WorkFilter(request.GET, queryset=cases)
        cards = filter.qs
        context = {'cards': cards, 'myFilter': filter}
        return render(request, 'work/cases.html', context)
    elif request.user.position.role_id == 3:
        cases = Case.objects.filter(user=request.user)
        country_cards = Case.objects.filter(country_id=request.user.country.country_id)
        filter = WorkFilter(request.GET, queryset=cases | country_cards)
        cards = filter.qs
        context = {'cards': cards, 'myFilter': filter}
        return render(request, 'work/cases.html', context)
    else:
        raise Http404('Недостаточно прав!')


def load_regions(request):
    country_id = request.GET.get('country_id')
    regions = Region.objects.filter(country=country_id).all()
    return render(request,'work/region_dropdown.html',{'regions':regions})

@login_required()
def add_comment(request, pk):

    case = Case.objects.get(id=pk)

    if request.method == 'POST':
        data = request.POST
        print(case.case_name)
        data._mutable = True
        data['case'] = case.id
        form = CaseCommentForm(data)
        if form.is_valid():
            from_email = request.user.email
            print(from_email)

            to_email = case.user.email
            print(to_email)
            comment = request.POST.get('comment')
            subject = "Новые комментарии"
            message = f"На вашу карточку /'{case.case_name}'/ был оставлен слудующий комментарий : '{comment}', oт {from_email}"
            form.save()
            send_mail(subject, message, settings.EMAIL_HOST_USER, [to_email, ], fail_silently=False)
            return redirect('works_list')
    else:
        form = CaseCommentForm()
    return render(request, 'work/add_comment.html', {'form': form, 'case': case})


def show_comments(request, pk):
    comments = CaseComment.objects.filter(case_id=pk, active=True)
    return render(request, 'work/show_comments.html', {'comments': comments})


def delete_comment(request, pk):
    CaseComment.objects.get(id=pk).delete()
    return redirect('works_list')


def case_render_pdf_view(request, *args, **kwargs):
    pk = kwargs.get('pk')
    case = get_object_or_404(Case, pk=pk)
    source = Source.objects.filter(case__pk=pk)
    intruder = Intruder.objects.filter(case__pk=pk)
    trade_union_activities = TradeUnionActivities.objects.filter(case__id=pk)
    comments = CaseComment.objects.filter(case_id=pk)
    template_path = 'work/work_pdf.html'
    context = {
        'case': case,
        'source': source,
        'intruder': intruder,
        'trade_union_activities': trade_union_activities,
        'comments': comments
    }
    # find the template and render it.
    template = get_template(template_path)
    html = template.render(context)
    # create a pdf
    pdf = pdfkit.from_string(html, False)
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'filename="card_{case.id}.pdf"'
    return response


def case_download_pdf_view(request, *args, **kwargs):
    pk = kwargs.get('pk')
    case = get_object_or_404(Case, pk=pk)
    source = Source.objects.filter(case__pk=pk)
    intruder = Intruder.objects.filter(case__pk=pk)
    trade_union_activities = TradeUnionActivities.objects.filter(case__id=pk)
    comments = CaseComment.objects.filter(case_id=pk)
    template_path = 'work/work_pdf.html'
    context = {
        'case': case,
        'source': source,
        'intruder': intruder,
        'trade_union_activities': trade_union_activities,
        'comments': comments
    }
    # find the template and render it.
    template = get_template(template_path)
    html = template.render(context)
    # create a pdf
    pdf = pdfkit.from_string(html, False)
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="card_{case.id}.pdf"'
    return response



class DataAPIView(APIView):
    def get(self, request):
        source = SourceSerializers(Source.objects.all(), many=True)
        country = CountrySerializers(Country.objects.all(), many=True)
        region_1 = RegionSerializers(Region.objects.filter(country__pk=1), many=True)
        region_2 = RegionSerializers(Region.objects.filter(country__pk=2), many=True)
        region_3 = RegionSerializers(Region.objects.filter(country__pk=3), many=True)
        region_4 = RegionSerializers(Region.objects.filter(country__pk=4), many=True)
        region_5 = RegionSerializers(Region.objects.filter(country__pk=5), many=True)
        groupOfRights = GroupOfRightsSerializers(GroupOfRights.objects.all(), many=True)
        tradeUnionRight = TradeUnionRightSerializers(TradeUnionRight.objects.all(), many=True)
        tradeUnionCrime = TradeUnionCrimeSerializers(TradeUnionCrime.objects.all(), many=True)
        meetingsRight = MeetingsRightSerializers(MeetingsRight.objects.all(), many=True)
        сonvention87 = Сonvention87Serializers(Сonvention87.objects.all(), many=True)
        tradeUnionBuildingsRight = TradeUnionBuildingsRightSerializers(TradeUnionBuildingsRight.objects.all(), many=True)
        createOrganizationRight = CreateOrganizationRightSerializers(CreateOrganizationRight.objects.all(), many=True)
        createTradeUnionRight  = CreateTradeUnionRightSerializers(CreateTradeUnionRight.objects.all(), many=True)
        electionsRight = ElectionsRightSerializers(ElectionsRight.objects.all(), many=True)
        tradeUnionActivityRight = TradeUnionActivityRightSerializers(TradeUnionActivityRight.objects.all(), many=True)
        createStrikeRight = CreateStrikeRightSerializers(CreateStrikeRight.objects.all(), many=True)
        сonvention98 = Сonvention98Serializers(Сonvention98.objects.all(), many=True)
        antiTradeUnionDiscrimination = AntiTradeUnionDiscriminationSerializers(AntiTradeUnionDiscrimination.objects.all(), many=True)
        conversationRight = СonversationRightSerializers(СonversationRight.objects.all(), many=True)
        сonvention135 = Сonvention135Serializers(Сonvention135.objects.all(), many=True)
        consultationRight = ConsultationRightSerializers(ConsultationRight.objects.all(), many=True)
        principleOfNonDiscrimination = PrincipleOfNonDiscriminationSerializers(PrincipleOfNonDiscrimination.objects.all(), many=True)
        discriminatiOnVariousGrounds = DiscriminatiOnVariousGroundsSerializers(DiscriminatiOnVariousGrounds.objects.all(), many=True)
        discriminationInVariousAreas = DiscriminationInVariousAreasSerializers(DiscriminationInVariousAreas.objects.all(), many=True)
        publicPolicyDiscrimination  = PublicPolicyDiscriminationSerializers(PublicPolicyDiscrimination.objects.all(), many=True)
        childLabor = ChildLaborSerializers(ChildLabor.objects.all(), many=True)
        сonvention138 = Сonvention138Serializers(Сonvention138.objects.all(), many=True)
        convention182 = Сonvention182Serializers(Сonvention182.objects.all(), many=True)
        prohibitionOfForcedLabor = ProhibitionOfForcedLaborSerializers(ProhibitionOfForcedLabor.objects.all(), many=True)
        useOfForcedLabor = UseOfForcedLaborSerializers(UseOfForcedLabor.objects.all(), many=True)
        governmentCoercion = GovernmentCoercionSerializers(GovernmentCoercion.objects.all(), many=True)
        violationsUsingCompulsoryLabor = ViolationsUsingCompulsoryLaborSerializer(ViolationsUsingCompulsoryLabor.objects.all(), many=True)
        failureSystemicMeasures = FailureSystemicMeasuresSerializers(FailureSystemicMeasures.objects.all(), many=True)

        victim = VictimSerializers(Victim.objects.all(), many=True)
        intruder = IntruderSerializers(Intruder.objects.all(), many=True)
        violation_nature = NatureViolationSerializers(NatureViolation.objects.all(), many=True)
        rights_state = RightsStateSerializers(RightsState.objects.all(), many=True)
        victim_situation = VictimSituationSerializers(VictimSituation.objects.all(), many=True)
        tradeUnionSituation = TradeUnionSituationSerializers(TradeUnionSituation.objects.all(), many=True)
        trade_union_activities = TradeUnionActivitiesSerializers(TradeUnionActivities.objects.all(), many=True)
        user = UserSerializers(User.objects.all(), many=True)

        # tradeUnionCount = TradeUnionCountSerializers(TradeUnionCount.objects.all(), many=True)
        # company = CompanySerializers(Company.objects.all(), many=True)
        # tradeUnionInfo = TradeUnionInfoSerializers(TradeUnionInfo.objects.all(), many=True)
        # groupOfPersons = GroupOfPersonsSerializers(GroupOfPersons.objects.all(), many=True)
        # russian_regions = Region.objects.filter()
        return Response([
            {'id': 'source', 'name': 'Источник информации', 'item': source.data},
            {'id': 'country', 'name': 'Страна', 'item': country.data},
            {'id': 'region', 'name': 'Регион', 'item': {1: region_1.data,
                                                        2: region_2.data,
                                                        3: region_3.data,
                                                        4: region_4.data,
                                                        5: region_5.data}},
            {'id': 'groupofrights', 'name': 'Группа прав', 'item': groupOfRights.data},
            {'id': 'tradeunionright', 'name': 'Нарушение в сфере профсоюзных прав и гражданских свобод',
             'item': tradeUnionRight.data},
            {'id': 'tradeunioncrime', 'name': 'Обвинения в преступном поведении в связи с профсоюзной деятельностью',
             'item': tradeUnionCrime.data},
            {'id': 'meetingsright', 'name': 'Нарушения права на проведение собраний и демонстраций',
             'item': meetingsRight.data},
            {'id': 'сonvention87', 'name': 'Нарушения положений Конвенции МОТ №87', 'item': сonvention87.data},
            {'id': 'tradeunionbuildingsright', 'name': 'Защита профсоюзных помещений и имущества профсоюзов',
             'item': tradeUnionBuildingsRight.data},
            {'id': 'createorganizationright', 'name': 'Создание организации без предварительного разрешения',
             'item': createOrganizationRight.data},
            {'id': 'createTradeUnionRight', 'name': 'Создание профсоюзов и вступление в профсоюзы',
             'item': createTradeUnionRight.data},
            {'id': 'electionsRight', 'name': 'Нарушение права свободно выбирать своих представителей',
             'item': electionsRight.data},
            {'id': 'tradeUnionActivityRight',
             'name': 'Нарушения права профсоюза организовывать деятельность своего аппарата',
             'item': tradeUnionActivityRight.data},
            {'id': 'createStrikeRight', 'name': 'Нарушение права на забастовку', 'item': createStrikeRight.data},
            {'id': 'сonvention98', 'name': 'Нарушения положений Конвенции МОТ №98', 'item': сonvention98.data},
            {'id': 'antiTradeUnionDiscrimination', 'name': 'Антипрофсоюзная дискриминация',
             'item': antiTradeUnionDiscrimination.data},
            {'id': 'conversationRight', 'name': 'Нарушения права на проведение коллективных переговоров',
             'item': conversationRight.data},
            {'id': 'сonvention135', 'name': 'Нарушения положений Конвенции МОТ №135', 'item': сonvention135.data},
            {'id': 'consultationRight', 'name': 'Проведение консультаций', 'item': consultationRight.data},
            {'id': 'principleOfNonDiscrimination', 'name': 'Принцип запрещения дискриминации',
             'item': principleOfNonDiscrimination.data},
            {'id': 'discriminatiOnVariousGrounds', 'name': 'Дискриминация по различным основаниям',
             'item': discriminatiOnVariousGrounds.data},
            {'id': 'discriminationInVariousAreas', 'name': 'Дискриминация в различных сферах трудовых отношений',
             'item': discriminationInVariousAreas.data},
            {'id': 'publicPolicyDiscrimination',
             'name': 'Нарушения в области проведения государственной политики по искоренению дискриминации и поощрению равенства прав и возможностей',
             'item': publicPolicyDiscrimination.data},
            {'id': 'childLabor', 'name': 'Дискриминация в различных сферах трудовых отношений',
             'item': childLabor.data},
            {'id': 'сonvention138', 'name': 'О минимальном возрасте для приема на работу', 'item': сonvention138.data},
            {'id': 'convention182',
             'name': 'О запрещении и немедленных мерах по искоренению наихудших форм детского труда',
             'item': convention182.data},
            {'id': 'prohibitionOfForcedLabor', 'name': 'Запрет принудительного труда',
             'item': prohibitionOfForcedLabor.data},
            {'id': 'useOfForcedLabor', 'name': 'Использование принудительного труда', 'item': useOfForcedLabor.data},
            {'id': 'governmentCoercion', 'name': 'Косвенное принуждение государством к труду',
             'item': governmentCoercion.data},
            {'id': 'violationsUsingCompulsoryLabor',
             'name': 'Нарушения при использовании принудительного (обязательного) труда в допустимых случаях',
             'item': violationsUsingCompulsoryLabor.data},
            {'id': 'failureSystemicMeasures', 'name': 'Нарушения, связанные с непринятием государством системных мер',
             'item': failureSystemicMeasures.data},

            {'id': 'victim', 'name': 'В отношении кого совершено нарушение', 'item': victim.data},
            {'id': 'intruder', 'name': 'Кем было совершено нарушение', 'item': intruder.data},
            {'id': 'natureviolation', 'name': 'Характер нарушения', 'item': violation_nature.data},
            {'id': 'rightsstate', 'name': 'Ситуация с правами', 'item': rights_state.data},
            {'id': 'victimsituation', 'name': 'Ситуация с потерпевшим(и)', 'item': victim_situation.data},
            {'id': 'tradeUnionSituation', 'name': 'Профсоюз на месте работы после произошедшего',
             'item': tradeUnionSituation.data},
            {'id': 'work_tradeunionactivities', 'name': 'Отрасль деятельности профсоюза', 'item': trade_union_activities.data},
            {'id': 'user', 'name': 'Монитор', 'item': user.data},
        ])
        # {'tradeUnionCount': {'Численность профсоюза после произошедшего': tradeUnionCount.data}}, #Можно удалить
        # {'company': {'Работодатель(компания)': company.data}}, #Можно удалить
        # {'tradeUnionInfo': {'Профсоюзная организация': tradeUnionInfo.data}}, #Можно удалить
        # {'groupOfPersons': {'Группа лиц (работников)': groupOfPersons.data}}, #Можно удалить


class DataFilterAPI(APIView):
    authentication_classes = []

    def post(self, request):

        print(request.data)
        one = get_request_data(request.data)
        two = return_right_data(one)
        values_list = get_values(two)
        dicts = get_data(two)
        print(dicts)

        queryset = Case.objects.filter(**dicts).values(*values_list).annotate(count=Count('id'),
                                                                                  procent=100 / Count('id')
                                                                                  )
        print(queryset)
        fields = get_fields(two)
        fields.append("count")
        fields.append("procent")
        serializers = DataFilterApiSerializer(queryset, many=True, fields=fields)
        return Response(serializers.data)


@login_required()
def case_files_download(request, pk):
    if request.user.position.role_id == 3:
        case = Case.objects.get(user=request.user, pk=pk)
        images = CasePhoto.objects.filter(card_id=case.id)
        files = CaseFile.objects.filter(card_id=case.id)
        return render(request, 'work/files_download.html', {'images': images, 'files': files})
    else:
        case = Case.objects.get(pk=pk)
        images = CasePhoto.objects.filter(card_id=case.id)
        files = CaseFile.objects.filter(card_id=case.id)
        return render(request, 'work/files_download.html', {'images': images, 'files': files})


@login_required()
def case_files_delete(request, pk):
    case_file = CaseFile.objects.get(pk=pk)
    case_id = case_file.card_id
    case_file.delete()
    return redirect('work_case_files_download', pk=case_id)


@login_required()
def case_photo_delete(request, pk):
    case_photo = CasePhoto.objects.get(pk=pk)
    case_id = case_photo.card_id
    case_photo.delete()
    return redirect('work_case_files_download', pk=case_id)

def work_pdf_in_html_page(request, pk):
    case = get_object_or_404(Case, pk=pk)
    source = Source.objects.filter(case__pk=pk)
    intruder = Intruder.objects.filter(case__pk=pk)
    comments = CaseComment.objects.filter(case_id=pk)
    trade_union_activities = TradeUnionActivities.objects.filter(case__id=pk)
    context = {
        'case': case,
        'source': source,
        'intruder': intruder,
        'trade_union_activities': trade_union_activities,
        'comments': comments
    }
    return render(request, 'work/work_pdf.html', context)


def work_word_generate(request, pk):
    case = get_object_or_404(Case, pk=pk)
    sources = Source.objects.filter(case__pk=pk)
    intruders = Intruder.objects.filter(case__pk=pk)
    comments = CaseComment.objects.filter(case_id=pk)
    trade_union_activities = TradeUnionActivities.objects.filter(case__id=pk)
    individual_infos = IndividualInfo.objects.filter(case__id=pk)
    document = Document()
    dates_list = ['date_create', 'date_update', 'start_date', 'end_date']
    document.add_heading('Трудовые нарушения', 0)

    records = []
    case_values = Case.objects.filter(pk=pk).values()
    field_name_list = [i.name for i in Case._meta.get_fields()]
    for field in field_name_list:
        try:
            if field == 'source':
                verbose_name = Case._meta.get_field(field).verbose_name
                for source in sources:
                    source_name = source.name
                    if source_name is not None and source_name != '':
                        records.append((verbose_name, source_name))
            elif field == 'intruder':
                verbose_name = Case._meta.get_field(field).verbose_name
                for intruder in intruders:
                    intruder_name = intruder.name
                    if intruder_name is not None and intruder_name != '':
                        records.append((verbose_name, intruder_name))
            elif field == 'trade_union_activities':
                verbose_name = Case._meta.get_field(field).verbose_name
                for trade_union_activitie in trade_union_activities:
                    trade_union_activitie_name = trade_union_activitie.name
                    if trade_union_activitie_name is not None and trade_union_activitie_name != '':
                        records.append((verbose_name, trade_union_activitie_name))
            elif field == 'individualinfo':
                verbose_name = IndividualInfo._meta.verbose_name
                for individual_info in individual_infos:
                    individual_info_name = individual_info.name
                    if individual_info_name is not None and individual_info_name != '':
                        records.append((verbose_name, individual_info_name))
            elif field == 'casephoto':
                continue
            elif field == 'casefile':
                continue
            elif field == 'casecomment':
                continue
            elif field == 'active':
                continue
            elif field == 'id':
                continue
            else:
                verbose_name = Case._meta.get_field(field).verbose_name
                value = case_values[0][field]
                if value is not None and value != '':
                    if field in dates_list:
                        records.append((verbose_name, value.strftime("%Y.%m.%d")))
                    else:
                        records.append((verbose_name, value))
        except KeyError:
            try:
                field += '_id'
                value_id = case_values[0][field]
                if value_id is not None and value_id != '':
                    value_name = Case._meta.get_field(field).related_model.objects.get(pk=value_id)
                    if field == 'tradeUnionCount_id':
                        records.append((verbose_name, value_name.choice))
                    elif field == 'tradeUnionInfo_id':
                        records.append((verbose_name, value_name.tradeunion_name))
                    elif field == 'groupOfPersons_id':
                        records.append((verbose_name, value_name.amount))
                    elif field == 'user_id':
                        records.append((verbose_name, value_name.username))
                    else:
                        records.append((verbose_name, value_name.name))
                else:
                    continue
            except Exception as e:
                print(e, field)

    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for field, value in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(field)
        row_cells[1].text = str(value)

    document.add_page_break()
    base_dir = str(settings.BASE_DIR)
    base_dir += "/work/static/word/work/"
    save_path = base_dir + f'card_{case.id}.docx'
    document.save(save_path)
    if os.path.exists(save_path):
        with open(save_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(save_path)
            return response
    return Http404()
