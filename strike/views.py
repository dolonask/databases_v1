import pprint

from django.contrib.auth.decorators import login_required
from django.core.mail import send_mail
from django.db.models import Count
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, Http404
from django.template.loader import get_template
from rest_framework.views import APIView
from rest_framework.response import Response
from main.service import unpucking
from .helpers import get_request_data, return_right_data, get_data, get_values, get_fields
from .serializers import *
from .filters import CardFilter
from django.db import connection
from .forms import *
from .models import *
from docxtpl import DocxTemplate
from migrant.templatetags.migrant_tags import var_verbose_name_for_word
from docx import Document
import jinja2
import os
import pdfkit



def form_formset(data):
    a = {}
    for i , y in data.items():
        if i.startswith('form'):
            a.update({i:y})
    for i in a.copy():
        # print(i.endswith('id'))
        if i.endswith("id"):
            a.pop(i)
    return a

def total_forms(data):

    if len(data['form-TOTAL_FORMS']) == 0:
        data['form-TOTAL_FORMS'] = '1'
        data['form-INITIAL_FORMS'] = '0'
        data['form-MAX_NUM_FORMS'] = '1000'
        data['form-MIN_NUM_FORMS'] = '0'
    # if len(data['form-TOTAL_FORMS']) == 2:
    #     data['form-TOTAL_FORMS'] = '2'
    #     data['form-INITIAL_FORMS'] = '0'
    #     data['form-MAX_NUM_FORMS'] = '1000'
    #     data['form-MIN_NUM_FORMS'] = '0'
    else:
        data['form-TOTAL_FORMS'] = data['form-TOTAL_FORMS']
        data['form-INITIAL_FORMS'] = '0'
        data['form-MAX_NUM_FORMS'] = '1000'
        data['form-MIN_NUM_FORMS'] = '0'
    return data

@login_required()
def add_case(request):
    general_tabs_fields = ['name',
                           'card_sources',
                           'source_url',
                           'source_content',
                           'country',
                           'region',
                           'city_name',
                           'company_name',
                           'company_ownership_type',
                           'company_is_tnk_member',
                           'company_tnk_name',
                           'company_employees_count',
                           'count_strike_participants',
                           'card_demand_categories',
                           'economic_demands',
                           'politic_demands',
                           'combo_demands',
                           'start_date',
                           'end_date',
                           'tradeunionChoice',
                           'tradeunionChoiceAnother',
                           'economic_another',
                           'politic_another',
                           'combo_another',
                           ]
    initiator_tab_fields = ['initiator']
    description_tab_fields = ['duration',
                              'meeting_requirements',
                              'story',
                              'reasons_for_strike',
                              'change_number_participants',
                              'initiators_and_participants',
                              'state_position',
                              'results_so_far',
                              'additional_information',

                              ]
    if request.method == "POST":
        form = CardForm(request.POST)
        pprint.pprint(request.POST)
        tradeUnionForm = TradeunionForm(request.POST)
        personGroupInfoForm = PersonGroupInfoForm(request.POST)
        # individualForm = IndividualForm(request.POST)
        data0 = form_formset(request.POST)
        data = total_forms(data0)
        individualFormSet = IndividualFormSet(data)

        employerForm = EmployerForm(request.POST)
        photoForm = CardPhotoForm(request.POST)
        fileForm = CardFileForm(request.POST)

        if form.is_valid():
            case = form.save(commit=False)

            if tradeUnionForm.is_valid():
                case.tradeunion_data = tradeUnionForm.save()
            if personGroupInfoForm.is_valid():
                case.personGroupInfo = personGroupInfoForm.save()
            if employerForm.is_valid():
                case.employear = employerForm.save()


            case.added_by = request.user
            case.active = True
            case.save()
            form._save_m2m()

            if individualFormSet.is_valid():
                individualFormSet.save(commit=False)



            for individual in individualFormSet:
                if individual.is_valid():
                    ind = individual.save(commit=False)
                    print(case.id)
                    ind.card_id = case.id
                    # print(ind._case)
                    ind.save()
            if photoForm.is_valid():
                for f in request.FILES.getlist('photo'):
                    photo = CardPhoto(photo=f, card=case)
                    photo.save()
            if fileForm.is_valid():
                for f in request.FILES.getlist('file'):
                    file = CardFile(file=f, card=case)
                    file.save()

            return redirect('strikes_list')
    else:
        form = CardForm
        tradeUnionForm = TradeunionForm
        personGroupInfoForm = PersonGroupInfoForm
        individualFormSet = IndividualFormSet(queryset=Individual.objects.none())
        employerForm = EmployerForm
        photoForm = CardPhotoForm
        fileForm = CardFileForm

    return render(request,'strike/add_case.html', context={
        'form':form,
        'tradeUnionForm':tradeUnionForm,
        'personGroupInfoForm':personGroupInfoForm,
        'individualFormSet':individualFormSet,
        'employerForm':employerForm,
        'photoForm':photoForm,
        'fileForm':fileForm,
        'general_tabs_fields':general_tabs_fields,
        'initiator_tab_fields':initiator_tab_fields,
        'description_tab_fields':description_tab_fields,
    })


def load_regions(request):
    country_id = request.GET.get('country_id')
    regions = Region.objects.filter(country=country_id).all()
    return render(request,'strike/region_dropdown.html',{'regions':regions})


@login_required()
def cases(request):
    if request.user.position.role_id == 1:
        cards = Card.objects.all()
        myFilter = CardFilter(request.GET, queryset=cards)
        cards = myFilter.qs
        context = {'cards': cards, 'myFilter': myFilter}
        return render(request, 'strike/strike.html', context)
    elif request.user.position.role_id == 2:
        cards = Card.objects.filter(country_id=request.user.country.country_id)
        myFilter = CardFilter(request.GET, queryset=cards)
        cards = myFilter.qs
        context = {'cards': cards, 'myFilter': myFilter}
        return render(request, 'strike/strike.html', context)
    elif request.user.position.role_id == 3:
        cards = Card.objects.filter(added_by=request.user)
        country_cards = Card.objects.filter(country_id=request.user.country.country_id)
        myFilter = CardFilter(request.GET, queryset=cards | country_cards)
        cards = myFilter.qs
        context = {'cards': cards, 'myFilter': myFilter}
        return render(request, 'strike/strike.html', context)
    else:
        raise Http404('Недостаточно прав!')


@login_required()
def delete_case(request, pk):
    case = Card.objects.get(id=pk).delete()
    return redirect('strikes_list')

@login_required()
def update_case(request, pk):
    case = Card.objects.get(id=pk)
    general_tabs_fields = ['name',
                           'card_sources',
                           'source_url',
                           'source_content',
                           'country',
                           'region',
                           'city_name',
                           'company_name',
                           'company_ownership_type',
                           'company_is_tnk_member',
                           'company_tnk_name',
                           'company_employees_count',
                           'count_strike_participants',
                           'card_demand_categories',
                           'economic_demands',
                           'politic_demands',
                           'combo_demands',
                           'start_date',
                           'end_date',
                           'tradeunionChoice',
                           'tradeunionChoiceAnother',
                           'economic_another',
                           'politic_another',
                           'combo_another',
                           ]
    initiator_tab_fields = ['initiator']
    description_tab_fields = ['duration',
                              'meeting_requirements',
                              'story',
                              'reasons_for_strike',
                              'change_number_participants',
                              'initiators_and_participants',
                              'state_position',
                              'results_so_far',
                              'additional_information',

                              ]
    if request.method == "POST":
        form = CardForm(request.POST, instance=case)
        tradeUnionForm = TradeunionForm(request.POST)
        personGroupInfoForm = PersonGroupInfoForm(request.POST)
        individualFormSet = IndividualFormSet(data=request.POST)


        employerForm = EmployerForm(request.POST)
        photoForm = CardPhotoForm(request.POST)
        fileForm = CardFileForm(request.POST)

        if form.is_valid():
            form.save(commit=False)

            if tradeUnionForm.is_valid():
                case.tradeunion_data = tradeUnionForm.save()
            if personGroupInfoForm.is_valid():
                case.personGroupInfo = personGroupInfoForm.save()

            for individual in individualFormSet:
                if individual.is_valid():
                    ind = individual.save(commit=False)
                    ind.card = case
            # individualFormSet.save()
            if employerForm.is_valid():
                case.employear = employerForm.save()


            case.added_by = request.user
            case.active = True
            case.save()

            if individualFormSet.is_valid():
                individualFormSet.save()

            form._save_m2m()

            if photoForm.is_valid():
                for f in request.FILES.getlist('photo'):
                    photo = CardPhoto(photo=f, card=case)
                    photo.save()
            if fileForm.is_valid():
                for f in request.FILES.getlist('file'):
                    file = CardFile(file=f, card=case)
                    file.save()

            return redirect('strikes_list')
    else:
        form = CardForm(instance=case)

        tradeUnionForm = TradeunionForm
        if case.tradeunion_data is not None:
            tradeUnionForm = TradeunionForm(instance=TradeunionData.objects.get(pk = case.tradeunion_data_id))

        personGroupInfoForm = PersonGroupInfoForm
        if case.personGroupInfo is not None:
            personGroupInfoForm = PersonGroupInfoForm(instance=PersonGroupInfo.objects.get(pk=case.personGroupInfo_id))

        employerForm = EmployerForm
        if case.employear is not None:
            employerForm = EmployerForm(instance=Employer.objects.get(pk=case.employear_id))

        photoForm = CardPhotoForm()
        fileForm = CardFileForm()
        images = CardPhoto.objects.filter(card_id=case.id)
        files = CardFile.objects.filter(card_id=case.id)
        individual_infos = Individual.objects.filter(card__id=pk)
        individualFormSet = IndividualFormSet(queryset=Individual.objects.filter(card_id=case.id))
        if len(individual_infos) != 0:
            individualFormSet.extra = 0
    return render(request, 'strike/add_case.html', context={
        'form': form,
        'tradeUnionForm': tradeUnionForm,
        'personGroupInfoForm': personGroupInfoForm,
        'individualFormSet': individualFormSet,
        'employerForm': employerForm,
        'photoForm': photoForm,
        'fileForm': fileForm,
        'general_tabs_fields': general_tabs_fields,
        'initiator_tab_fields': initiator_tab_fields,
        'description_tab_fields': description_tab_fields,
        'images': images,
        'files': files,
    })


@login_required()
def add_comment(request, pk):
    card = Card.objects.get(id=pk)
    if request.method == 'POST':
        data = request.POST
        data._mutable = True
        data['card'] = card.id
        form = CardCommentForm(data)
        if form.is_valid():
            from_email = request.user.email
            to_email = card.added_by.email
            comment = request.POST.get('comment')
            subject = "Новые комментарии"
            message = f"На вашу карточку /'{card.name}'/ был оставлен слудующий комментарий : '{comment}', oт {from_email}"
            form.save()
            send_mail(subject, message, settings.EMAIL_HOST_USER, [to_email, ], fail_silently=False)
            return redirect('strikes_list')
    else:
        form = CardCommentForm()
    return render(request, 'strike/add_comment.html', {'form': form, 'card': card})


def show_comments(request, pk):
    comments = CardComment.objects.filter(card_id=pk, active=True)
    return render(request, 'strike/show_comments.html', {'comments': comments})


def delete_comment(request, pk):
    CardComment.objects.get(id=pk).delete()
    return redirect('strikes_list')


def case_render_pdf_view(request, *args, **kwargs):
    pk = kwargs.get('pk')
    card = get_object_or_404(Card, pk=pk)
    card_sources = Source.objects.filter(source__pk=pk)
    card_demand_categories = DemandCategory.objects.filter(card__pk=pk)
    economic_demands = EconomicDemand.objects.filter(card__pk=pk)
    politic_demands = PoliticDemand.objects.filter(card__pk=pk)
    combo_demands = ComboDemand.objects.filter(card__pk=pk)
    comments = CardComment.objects.filter(card_id=pk)
    template_path = 'strike/strike_pdf.html'
    context = {
        'card': card,
        'card_sources': card_sources,
        'card_demand_categories': card_demand_categories,
        'economic_demands': economic_demands,
        'politic_demands': politic_demands,
        'combo_demands': combo_demands,
        'comments': comments,
    }
    # find the template and render it.
    template = get_template(template_path)
    html = template.render(context)
    # create a pdf
    pdf = pdfkit.from_string(html, False)
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'filename="card_{card.id}.pdf"'
    # StringIO(html.encode("UTF-8")), response, encoding='UTF-8')
    # if error then show some funy view
    return response


def case_download_pdf_view(request, *args, **kwargs):
    pk = kwargs.get('pk')
    card = get_object_or_404(Card, pk=pk)
    card_sources = Source.objects.filter(card__pk=pk)
    card_demand_categories = DemandCategory.objects.filter(card__pk=pk)
    economic_demands = EconomicDemand.objects.filter(card__pk=pk)
    politic_demands = PoliticDemand.objects.filter(card__pk=pk)
    combo_demands = ComboDemand.objects.filter(card__pk=pk)
    comments = CardComment.objects.filter(card_id=pk)
    template_path = 'strike/strike_pdf.html'
    context = {
        'card': card,
        'card_sources': card_sources,
        'card_demand_categories': card_demand_categories,
        'economic_demands': economic_demands,
        'politic_demands': politic_demands,
        'combo_demands': combo_demands,
        'comments': comments,
    }
    # Create a Django response object, and specify content_type as pdf
    # find the template and render it.
    template = get_template(template_path)
    html = template.render(context)
    # create a pdf
    pdf = pdfkit.from_string(html, False)
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="card_{card.id}.pdf"'
    return response

class DataAPIView(APIView):
    def get(self, request):
        country = CountrySerializers(Country.objects.all(), many=True)
        region_1 = RegionSerializers(Region.objects.filter(country__pk=1), many=True)
        region_2 = RegionSerializers(Region.objects.filter(country__pk=2), many=True)
        region_3 = RegionSerializers(Region.objects.filter(country__pk=3), many=True)
        region_4 = RegionSerializers(Region.objects.filter(country__pk=4), many=True)
        region_5 = RegionSerializers(Region.objects.filter(country__pk=5), many=True)
        source = SourceSerializers(Source.objects.all(), many=True)
        added_by = UserSerializers(User.objects.all(), many=True)
        demand_categories = DemandCategorySerializers(DemandCategory.objects.all(), many=True)
        economic_demands = EconomicDemandSerializers(EconomicDemand.objects.all(), many=True)
        politic_demands = PoliticDemandSerializers(PoliticDemand.objects.all(), many=True)
        combo_demands = ComboDemandSerializers(ComboDemand.objects.all(), many=True)
        company_ownership_type = OwnerShipTypeSerializers(OwnerShipType.objects.all(), many=True)
        company_employees_count = EmployeesCountSerializers(EmployeesCount.objects.all(), many=True)
        count_strike_participants = ParticipantsCountSerializers(ParticipantsCount.objects.all(), many=True)
        tradeunion_choice = TradeunionChoiceSerializers(TradeunionChoice.objects.all(), many=True)
        initiator = InitiatorSerializers(Initiator.objects.all(), many=True)
        tradeunion_data = TradeunionDataSerializers(TradeunionData.objects.all(), many=True)
        employear = EmployerSerializers(Employer.objects.all(), many=True)
        duration = StrikeCharacterSerializers(StrikeCharacter.objects.all(), many=True)
        meeting_requirements = MeetingRequirmentSerializers(MeetingRequirment.objects.all(), many=True)
        # personGroupInfo = Группа лиц
        return Response([
            {'id': 'country', 'name': 'Страна', 'item': country.data},
            {'id': 'region', 'name': 'Регион', 'item': {1: region_1.data,
                                                        2: region_2.data,
                                                        3: region_3.data,
                                                        4: region_4.data,
                                                        5: region_5.data}},
            {'id': 'card_sources', 'name': 'Источник', 'item': source.data},
            {'id': 'added_by', 'name': 'Монитор', 'item': added_by.data},
            {'id': 'card_demand_categories', 'name': 'Характер требований', 'item': demand_categories.data},
            {'id': 'economic_demands', 'name': 'Экономический', 'item': economic_demands.data},
            {'id': 'politic_demands', 'name': 'Политический', 'item': politic_demands.data},
            {'id': 'combo_demands', 'name': 'Смешанный', 'item': combo_demands.data},
            {'id': 'company_ownership_type', 'name': 'Форма собственности компании',
             'item': company_ownership_type.data},

            {'id': 'company_employees_count', 'name': 'Общая численность работников на предприятии',
             'item': company_employees_count.data},

            {'id': 'count_strike_participants', 'name': 'Общая численность работников на предприятии',
             'item': count_strike_participants.data},

            {'id': 'tradeunionChoice', 'name': 'Есть ли на предприятии профсоюз', 'item': tradeunion_choice.data},
            {'id': 'initiator', 'name': 'Инициатор забастовки/акции', 'item': initiator.data},
            {'id': 'tradeunion_data', 'name': 'Данные профсоюза', 'item': tradeunion_data.data},
            {'id': 'employear', 'name': 'Работодатель', 'item': employear.data},
            {'id': 'duration', 'name': 'Характер забастовки/акции - сколько длилась', 'item': duration.data},
            {'id': 'meeting_requirements', 'name': 'Удовлетворение требований', 'item': meeting_requirements.data},
        ])

class TestGet(APIView):
    def get(self, request):

        queryset = Source.objects.filter(id=4).values("id", "name")
        print(queryset)
        serializer = TestSerializer(queryset, many=True)
        # print(serializer)
        return Response(serializer.data)

class DataFilterAPI(APIView):
    authentication_classes = []
    def post(self, request):
        print(request.data)
        one = get_request_data(request.data)
        two = return_right_data(one)
        values_list = get_values(two)
        dicts = get_data(two)
        print(values_list, "values")
        # print(dicts)

        queryset = Card.objects.filter(**dicts).values(*values_list).annotate(count=Count('id'),
                                                                         procent=100 / Count('id'))
        print(queryset)

        fields = get_fields(two)
        print(fields, "f")
        fields.append("count")
        fields.append("procent")

        serializers = FilterStrikeGroupSerializer(queryset, many=True, fields=fields)
        return Response(serializers.data)


@login_required()
def card_files_download(request, pk):
    if request.user.position.role_id == 3:
        case = Card.objects.get(added_by=request.user, id=pk)
        images = CardPhoto.objects.filter(card_id=case.id)
        files = CardFile.objects.filter(card_id=case.id)
        return render(request, 'strike/files_download.html', {'images': images, 'files': files})
    else:
        case = Card.objects.get(id=pk)
        images = CardPhoto.objects.filter(card_id=case.id)
        files = CardFile.objects.filter(card_id=case.id)
        return render(request, 'strike/files_download.html', {'images': images, 'files': files})


@login_required()
def card_files_delete(request, pk):
    card_file = CardFile.objects.get(pk=pk)
    card_id = card_file.card_id
    card_file.delete()
    return redirect('strike_card_files_download', pk=card_id)


@login_required()
def card_photo_delete(request, pk):
    card_photo = CardPhoto.objects.get(pk=pk)
    card_id = card_photo.card_id
    card_photo.delete()
    return redirect('strike_card_files_download', pk=card_id)

def strike_word_generate(request, pk):
    card = Card.objects.get(id=pk)
    card_sources = Source.objects.filter(card__pk=pk)
    card_demand_categories = DemandCategory.objects.filter(card__pk=pk)
    economic_demands = EconomicDemand.objects.filter(card__pk=pk)
    politic_demands = PoliticDemand.objects.filter(card__pk=pk)
    combo_demands = ComboDemand.objects.filter(card__pk=pk)

    document = Document()
    dates_list = ['date_create', 'date_update', 'start_date', 'end_date']
    document.add_heading('Забастовка', 0)
    records = []
    card_values = Card.objects.filter(pk=pk).values()
    field_name_list = [i.name for i in Card._meta.get_fields()]
    for field in field_name_list:
        try:
            if field == 'card_sources':
                verbose_name = Card._meta.get_field(field).verbose_name
                for source in card_sources:
                    source_name = source.name
                    if source_name is not None and source_name != '':
                        records.append((verbose_name, source_name))
            elif field == 'card_demand_categories':
                verbose_name = Card._meta.get_field(field).verbose_name
                for card_demand_category in card_demand_categories:
                    card_demand_category_name = card_demand_category.demand_cat_name
                    if card_demand_category_name is not None and card_demand_category_name != '':
                        records.append((verbose_name, card_demand_category_name))
            elif field == 'economic_demands':
                verbose_name = Card._meta.get_field(field).verbose_name
                for economic_demand in economic_demands:
                    economic_demand_name = economic_demand.name
                    if economic_demand_name is not None and economic_demand_name != '':
                        records.append((verbose_name, economic_demand_name))
            elif field == 'politic_demands':
                verbose_name = Card._meta.get_field(field).verbose_name
                for politic_demand in politic_demands:
                    politic_demand_name = politic_demand.name
                    if politic_demand_name is not None and politic_demand_name != '':
                        records.append((verbose_name, politic_demand_name))
            elif field == 'combo_demands':
                verbose_name = Card._meta.get_field(field).verbose_name
                for combo_demand in combo_demands:
                    combo_demand_name = combo_demand.name
                    if combo_demand_name is not None and combo_demand_name != '':
                        records.append((verbose_name, combo_demand_name))
            elif field == 'individual':
                continue
            elif field == 'cardphoto':
                continue
            elif field == 'cardfile':
                continue
            elif field == 'cardcomment':
                continue
            elif field == 'active':
                continue
            elif field == 'id':
                continue
            else:
                verbose_name = Card._meta.get_field(field).verbose_name
                value = card_values[0][field]
                if value is not None and value != '':
                    if field in dates_list:
                        records.append((verbose_name, value.strftime("%Y.%m.%d")))
                    else:
                        records.append((verbose_name, value))
        except KeyError:
            try:
                field += '_id'
                value_id = card_values[0][field]
                if value_id is not None and value_id != '':
                    value_name = Card._meta.get_field(field).related_model.objects.get(pk=value_id)
                    if field == 'company_employees_count_id' or field == 'count_strike_participants_id':
                        records.append((verbose_name, value_name.choice))
                    elif field == 'tradeunion_data_id':
                        records.append((verbose_name, value_name.tradeUnion_name))
                    elif field == 'added_by_id':
                        records.append((verbose_name, value_name.username))
                    else:
                        records.append((verbose_name, value_name.name))
                else:
                    continue
            except Exception as e:
                print(e, field)

    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    # hdr_cells = table.rows[0].cells
    for field, value in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(field)
        row_cells[1].text = str(value)

    document.add_page_break()
    base_dir = str(settings.BASE_DIR)
    base_dir += "/strike/static/word/strike/"
    save_path = base_dir + f'card_{card.id}.docx'
    document.save(save_path)
    if os.path.exists(save_path):
        with open(save_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(save_path)
            return response
    return Http404()